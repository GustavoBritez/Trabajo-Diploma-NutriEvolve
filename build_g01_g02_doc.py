import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    primary_color = RGBColor(46, 117, 89)   # Dark Sage / Medical Green (#2E7559)
    secondary_color = RGBColor(52, 73, 94)  # Charcoal Slate (#34495E)
    text_color = RGBColor(44, 62, 80)       # Dark Text (#2C3E50)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TRABAJO DE DIPLOMA - INGENIERÍA DE SISTEMAS INFORMATICOS")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = secondary_color

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("PROYECTO: NutriEvolve - Plataforma de Monitoreo Nutricional Pediátrico y Análisis Demográfico OMS")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = primary_color

    p_doc_type = doc.add_paragraph()
    p_doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dt = p_doc_type.add_run("Documento de Especificación de Requisitos (G01. Propósito y G02. Alcance Pediátrico)")
    run_dt.font.name = 'Calibri'
    run_dt.font.size = Pt(12)
    run_dt.font.italic = True
    run_dt.font.color.rgb = secondary_color

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Header Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Universidad / Facultad:", "Universidad Abierta Interamericana - Facultad de Tecnología Informática"),
        ("Materia / Cursada:", "Ingeniería de Software / Trabajo de Diploma de Ingeniería"),
        ("Alumnos Integrantes:", "Britez Alvarenga Gustavo Daniel | Britez Alvarenga Leandro Ezequiel | Palavecino Fernando Daniel"),
        ("Docente / Sistema:", "Ing. Jorge Agustín Pereyra | Sistema NutriEvolve (Versión Pediátrica OMS)")
    ]

    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.name = 'Calibri'
        r_lbl.font.color.rgb = secondary_color

        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.name = 'Calibri'

        set_cell_background(cell_lbl, "F2F4F4")
        set_cell_background(cell_val, "FAFAFA")
        set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Section G01
    h_g01 = doc.add_paragraph()
    r_g01 = h_g01.add_run("G01. Propósito")
    r_g01.font.name = 'Calibri'
    r_g01.font.size = Pt(16)
    r_g01.font.bold = True
    r_g01.font.color.rgb = primary_color
    h_g01.paragraph_format.space_after = Pt(8)

    p_g01_1 = doc.add_paragraph()
    p_g01_1.paragraph_format.line_spacing = 1.15
    p_g01_1.paragraph_format.space_after = Pt(8)
    r = p_g01_1.add_run("En la práctica clínica actual de la nutrición pediátrica, tanto en consultorios independientes como en PyMEs del sector de la salud, el monitoreo del crecimiento infantil enfrenta deficiencias operativas y tecnológicas severas. La gestión de agendas de atención se encuentra habitualmente desvinculada del seguimiento clínico antropométrico, generando una dinámica administrativa fragmentada entre el área de recepción y el equipo de nutricionistas.")
    r.font.name = 'Calibri'; r.font.size = Pt(11)

    p_g01_2 = doc.add_paragraph()
    p_g01_2.paragraph_format.line_spacing = 1.15
    p_g01_2.paragraph_format.space_after = Pt(8)
    r = p_g01_2.add_run("Actualmente, la evaluación del desarrollo físico en niños de 0 a 5 años (y hasta 19 años) se basa en la aplicación manual de los Patrones de Crecimiento Infantil de la Organización Mundial de la Salud (OMS). Los profesionales suelen volcar manualmente mediciones de peso, talla y perímetro cefálico sobre curvas impresas en papel o archivos PDF estáticos. Este proceso tradicional presenta marcadas limitaciones:\n")
    r.font.name = 'Calibri'; r.font.size = Pt(11)

    bullet1 = doc.add_paragraph(style='List Bullet')
    r = bullet1.add_run("Alto consumo de tiempo en consulta: ")
    r.bold = True
    bullet1.add_run("El trazado y la interpolación visual punto a punto sobre las gráficas de percentiles restan tiempo valioso para la atención clínica directa del paciente y la orientación a sus tutores.")
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    r = bullet2.add_run("Riesgo de error humano en el diagnóstico: ")
    r.bold = True
    bullet2.add_run("La estimación 'a ojo' de percentiles intermedios o desvíos estándar impide calcular con precisión matemática los Z-Scores oficializados por la OMS, lo que incrementa el margen de error en la detección oportuna de condiciones críticas como la emaciación (desnutrición aguda), el desmedro (retraso estatural) o el riesgo de sobrepeso.")

    bullet3 = doc.add_paragraph(style='List Bullet')
    r = bullet3.add_run("Subutilización y fragmentación de la información (Pérdida de valor demográfico): ")
    r.bold = True
    bullet3.add_run("Los datos antropométricos registrados quedan aislados en historias clínicas individuales de papel o planillas locales. Esto imposibilita a los consultorios recopilar y consolidar datos epidemiológicos clave. Por ejemplo, en un consultorio ubicado en la localidad de Lanús (Buenos Aires), no es posible responder preguntas como: 'De 10.000 niños atendidos en esta región, ¿qué porcentaje presenta bajo peso o sobrepeso y qué planes alimentarios fueron prescritos?'.")

    p_g01_3 = doc.add_paragraph()
    p_g01_3.paragraph_format.line_spacing = 1.15
    p_g01_3.paragraph_format.space_before = Pt(8)
    p_g01_3.paragraph_format.space_after = Pt(8)
    r = p_g01_3.add_run("Como respuesta a la problemática planteada, el presente Trabajo de Diploma propone el desarrollo de ")
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    r_bold = p_g01_3.add_run("NutriEvolve")
    r_bold.bold = True; r_bold.font.name = 'Calibri'; r_bold.font.size = Pt(11)
    r_2 = p_g01_3.add_run(", una plataforma tecnológica integral diseñada específicamente para optimizar la gestión de turnos médicos y automatizar el monitoreo clínico pediátrico bajo normas OMS, incorporando un potente módulo de Inteligencia de Datos Demográficos.")
    r_2.font.name = 'Calibri'; r_2.font.size = Pt(11)

    p_g01_4 = doc.add_paragraph()
    p_g01_4.paragraph_format.line_spacing = 1.15
    p_g01_4.paragraph_format.space_after = Pt(14)
    r = p_g01_4.add_run("El sistema permitirá registrar controles trimestrales de mediciones antropométricas, graficar dinámicamente la evolución temporal del niño sobre las bandas de percentiles estándar (P3, P15, P50, P85, P97) y diagnosticar automáticamente su estado nutricional. Asimismo, NutriEvolve centralizará esta información de manera segura y anonimizada para generar reportes estadísticos demográficos regionales, permitiendo al profesional o a la institución de salud comprender la prevalencia nutricional y las prescripciones asociadas según la localización geográfica de la población atendida.")
    r.font.name = 'Calibri'; r.font.size = Pt(11)

    # Section G02
    h_g02 = doc.add_paragraph()
    r_g02 = h_g02.add_run("G02. Descripción funcional del producto y Alcance")
    r_g02.font.name = 'Calibri'
    r_g02.font.size = Pt(16)
    r_g02.font.bold = True
    r_g02.font.color.rgb = primary_color
    h_g02.paragraph_format.space_after = Pt(8)

    p_g02_intro = doc.add_paragraph()
    p_g02_intro.paragraph_format.line_spacing = 1.15
    p_g02_intro.paragraph_format.space_after = Pt(10)
    r = p_g02_intro.add_run("A continuación se detallan los requisitos funcionales del sistema, divididos en tres módulos operativos principales: Gestión de Turnos Médicos, Seguimiento Nutricional Pediátrico con Evaluación OMS, y Análisis Estadístico Demográfico Poblacional.")
    r.font.name = 'Calibri'; r.font.size = Pt(11)

    # RFN 1 Subtitle
    h_rfn1 = doc.add_paragraph()
    r_rfn1 = h_rfn1.add_run("RFN 1 : Gestión de Turnos Médicos y Agenda de Consultorio")
    r_rfn1.font.name = 'Calibri'; r_rfn1.font.size = Pt(13); r_rfn1.font.bold = True
    r_rfn1.font.color.rgb = secondary_color
    h_rfn1.paragraph_format.space_after = Pt(6)

    # RFN1 Table
    rfn1_table = doc.add_table(rows=1, cols=4)
    rfn1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = rfn1_table.rows[0].cells
    hdr_titles = ["ID", "Requisito", "Prioridad", "Descripción"]
    widths = [Inches(0.8), Inches(1.8), Inches(0.9), Inches(3.0)]

    for idx, title in enumerate(hdr_titles):
        cell = hdr_cells[idx]
        cell.width = widths[idx]
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "2E7559")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    rfn1_data = [
        ("RFN1.1", "Registrar Turno Médico", "Alta", "El sistema deberá permitir registrar un turno médico almacenando el código de turno, paciente (tutor y niño), profesional, fecha, horario, motivo de consulta, estado y observaciones."),
        ("RFN1.2", "Consultar Disponibilidad Profesional", "Alta", "El sistema deberá mostrar la disponibilidad de los nutricionistas para una fecha determinada, excluyendo los horarios ocupados por turnos previamente asignados."),
        ("RFN1.3", "Validar Superposición de Turnos", "Alta", "El sistema deberá impedir el registro o modificación de un turno cuando exista otro turno asignado al mismo profesional en la misma fecha y horario."),
        ("RFN1.4", "Actualizar Agenda Profesional", "Alta", "El sistema deberá actualizar automáticamente la agenda del profesional marcando como ocupado o disponible el horario afectado por el registro, modificación o cancelación de un turno."),
        ("RFN1.5", "Modificar Turno Médico", "Media", "El sistema deberá permitir modificar la fecha, horario y profesional asignado de un turno previamente registrado."),
        ("RFN1.6", "Buscar Turno Médico", "Media", "El sistema deberá permitir buscar turnos utilizando el DNI del tutor/paciente, el código de turno o la fecha de consulta."),
        ("RFN1.7", "Cancelar Turno Médico", "Alta", "El sistema deberá permitir cancelar un turno registrando el motivo de cancelación, la fecha de cancelación y el estado correspondiente."),
        ("RFN1.8", "Liberar Horario Cancelado", "Media", "Al cancelar un turno, el sistema deberá marcar como disponible el horario correspondiente en la agenda del profesional."),
        ("RFN1.9", "Registrar Estado del Turno", "Alta", "El sistema deberá permitir registrar el estado final del turno con uno de los siguientes valores: Asistió, Ausente o Cancelado.")
    ]

    for row_idx, data in enumerate(rfn1_data):
        row = rfn1_table.add_row()
        bg = "F9FAFA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
            if col_idx == 0: r.font.bold = True
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # RFN 2 Subtitle
    h_rfn2 = doc.add_paragraph()
    r_rfn2 = h_rfn2.add_run("RFN 2 : Seguimiento Nutricional Pediátrico y Curvas OMS")
    r_rfn2.font.name = 'Calibri'; r_rfn2.font.size = Pt(13); r_rfn2.font.bold = True
    r_rfn2.font.color.rgb = secondary_color
    h_rfn2.paragraph_format.space_after = Pt(6)

    # RFN2 Table
    rfn2_table = doc.add_table(rows=1, cols=4)
    rfn2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells2 = rfn2_table.rows[0].cells

    for idx, title in enumerate(hdr_titles):
        cell = hdr_cells2[idx]
        cell.width = widths[idx]
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "2E7559")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    rfn2_data = [
        ("RFN2.1", "Buscar Paciente Pediátrico", "Alta", "El sistema localiza un paciente infantil mediante DNI del tutor, DNI del niño, nombre o apellido para acceder a su historia clínica pediátrica unificada."),
        ("RFN2.2", "Registrar Mediciones Antropométricas Pediátricas", "Alta", "El sistema registra las mediciones del niño tomadas periódicamente (cada 3 meses o según consulta), incluyendo peso (kg), talla/longitud (cm), perímetro cefálico (cm) y cálculo automático de IMC."),
        ("RFN2.3", "Generar Curvas de Crecimiento OMS y Percentiles", "Alta", "El sistema grafica automáticamente la evolución del niño sobre las curvas patrón de la OMS (Peso para la Edad, Talla para la Edad, IMC para la Edad) visualizando las bandas P3, P15, P50, P85 y P97."),
        ("RFN2.4", "Calcular Z-Score y Diagnóstico Automatizado OMS", "Alta", "El sistema aplica el algoritmo LMS de la OMS para determinar el Z-Score del niño y emite un diagnóstico nutricional automatizado (Desnutrición Severa, Desnutrición, Normopeso, Riesgo Sobrepeso, Sobrepeso, Obesidad)."),
        ("RFN2.5", "Registrar Anamnesis Pediátrica y Antecedentes", "Alta", "El sistema registra datos del desarrollo infantil, tipo de lactancia (materna/fórmula), introducción de alimentación complementaria, alergias y antecedentes clínicos de la familia."),
        ("RFN2.6", "Registrar Evaluación Dietaria Pediátrica", "Alta", "El sistema registra la ingesta habitual del niño mediante el Recordatorio de 24 Horas y frecuencia de consumo de grupos de alimentos adaptados a la etapa infantil."),
        ("RFN2.7", "Prescribir Plan Alimentario Pediátrico", "Alta", "El sistema registra un plan alimentario personalizado según la edad en meses/años, indicando requerimientos calóricos, distribución de macronutrientes y pautas nutricionales para la familia."),
        ("RFN2.8", "Generar Alertas Clínicas de Crecimiento", "Media", "El sistema genera alertas visuales automáticas en caso de caídas abruptas de percentil entre controles consecutivos o desviaciones críticas de la mediana P50."),
        ("RFN2.9", "Consultar Historial Clínico Evolutivo Pediátrico", "Alta", "El sistema presenta de forma centralizada todas las consultas anteriores, evoluciones antropométricas, gráficos históricos y observaciones registradas para el paciente."),
        ("RFN2.10", "Actualizar Historial y Recalcular Gráficos", "Alta", "Al incorporar una nueva consulta de control, el sistema actualiza de manera inmediata la historia clínica del niño y recalcula las curvas evolutivas."),
        ("RFN2.11", "Registrar Objetivos del Tratamiento Pediátrico", "Media", "El sistema registra metas de salud infantil, incluyendo recomendaciones de hidratación, pautas de crianza alimentaria respetuosa y actividad física según la edad.")
    ]

    for row_idx, data in enumerate(rfn2_data):
        row = rfn2_table.add_row()
        bg = "F9FAFA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
            if col_idx == 0: r.font.bold = True
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # RFN 3 Subtitle
    h_rfn3 = doc.add_paragraph()
    r_rfn3 = h_rfn3.add_run("RFN 3 : Inteligencia y Estadística Demográfica Poblacional")
    r_rfn3.font.name = 'Calibri'; r_rfn3.font.size = Pt(13); r_rfn3.font.bold = True
    r_rfn3.font.color.rgb = secondary_color
    h_rfn3.paragraph_format.space_after = Pt(6)

    # RFN3 Table
    rfn3_table = doc.add_table(rows=1, cols=4)
    rfn3_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells3 = rfn3_table.rows[0].cells

    for idx, title in enumerate(hdr_titles):
        cell = hdr_cells3[idx]
        cell.width = widths[idx]
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "2E7559")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    rfn3_data = [
        ("RFN3.1", "Consolidar Datos Demográficos Anonimizados", "Alta", "El sistema agrupa los registros antropométricos atendidos en el consultorio/clínica resguardando la privacidad del paciente y asociándolos a variables de ubicación geográfica (Localidad, Partido, Provincia)."),
        ("RFN3.2", "Generar Reportes Epidemiológicos por Localidad", "Alta", "El sistema permite consultar indicadores poblacionales filtrados por región (ej. Lanús, Avellaneda, CABA), mostrando la cantidad total de niños atendidos y distribuciones porcentuales por diagnóstico OMS."),
        ("RFN3.3", "Analizar Prevalencia Nutricional Poblacional", "Alta", "El sistema emite estadísticas demográficas sobre la proporción de la población infantil atendida con normopeso, desnutrición, sobrepeso u obesidad en un rango de fechas determinado."),
        ("RFN3.4", "Correlacionar Diagnósticos y Prescripciones Alimentarias", "Media", "El sistema proporciona reportes comparativos indicando qué planes alimentarios o suplementaciones fueron prescritos según el diagnóstico nutricional y la franja etaria de la localidad."),
        ("RFN3.5", "Exportar Informes Demográficos Estadísticos", "Media", "El sistema permite exportar los tableros e informes demográficos a formatos PDF o Excel para la toma de decisiones clínicas y la presentación institucional.")
    ]

    for row_idx, data in enumerate(rfn3_data):
        row = rfn3_table.add_row()
        bg = "F9FAFA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
            if col_idx == 0: r.font.bold = True
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section G03 Definitions
    h_g03 = doc.add_paragraph()
    r_g03 = h_g03.add_run("G03. Definiciones, Acrónimos y Abreviaturas")
    r_g03.font.name = 'Calibri'
    r_g03.font.size = Pt(16)
    r_g03.font.bold = True
    r_g03.font.color.rgb = primary_color
    h_g03.paragraph_format.space_after = Pt(8)

    # Definitions Table
    def_table = doc.add_table(rows=1, cols=2)
    def_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    def_table.rows[0].cells[0].width = Inches(2.2)
    def_table.rows[0].cells[1].width = Inches(4.3)

    for idx, t in enumerate(["Término / Acrónimo", "Definición Conceptual en NutriEvolve"]):
        c = def_table.rows[0].cells[idx]
        p = c.paragraphs[0]; r = p.add_run(t)
        r.font.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "34495E")
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    defs = [
        ("Percentiles OMS", "Valores estadísticos de referencia que indican el porcentaje de la población infantil sana que se encuentra por debajo de una medición dada. Las curvas estándar muestran las líneas P3, P15, P50 (mediana), P85 y P97."),
        ("Z-Score (Puntuación Z)", "Unidad de medida estadística que expresa la distancia de una medición individual con respecto a la mediana de la población de referencia de la OMS, medida en desviaciones estándar."),
        ("Método LMS (OMS)", "Método matemático desarrollado por Cole y Green adoptado por la OMS, caracterizado por tres parámetros: Box-Cox power (L), Mediana (M) y Coeficiente de Variación (S), utilizado para calcular Z-scores continuos según la edad en meses/días y sexo."),
        ("Emaciación (Bajo Peso/Talla)", "Indicador de desnutrición aguda caracterizado por un peso sustancialmente inferior al correspondiente a la talla del niño."),
        ("Desmedro (Bajo Talla/Edad)", "Indicador de desnutrición crónica caracterizado por un retraso en el crecimiento estatural del niño respecto a su edad."),
        ("Antropometría Pediátrica", "Conjunto de mediciones físicas realizadas al cuerpo del niño (peso, longitud/talla, perímetro cefálico) para evaluar su crecimiento y estado de salud corporal."),
        ("Estadística Demográfica", "Módulo de análisis poblacional que agrupa datos nutricionales anonimizados por variables geográficas (localidad, partido) para emitir métricas epidemiológicas institucionales."),
        ("DVH / DVV", "Dígitos Verificadores Horizontales y Verticales. Mecanismo criptográfico interno del sistema para garantizar la integridad de las tablas de datos clínicos y evitar modificaciones no autorizadas en SQL Server."),
        ("Bitácora de Auditoría", "Registro automatizado e inalterable de todos los eventos del sistema, clasificándolos por nivel de criticidad para auditoría técnica y seguridad.")
    ]

    for row_idx, (term, dfn) in enumerate(defs):
        row = def_table.add_row()
        bg = "F9FAFA" if row_idx % 2 == 1 else "FFFFFF"
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = Inches(2.2), Inches(4.3)
        
        p0 = c0.paragraphs[0]; r0 = p0.add_run(term)
        r0.font.bold = True; r0.font.size = Pt(9.5); r0.font.name = 'Calibri'
        
        p1 = c1.paragraphs[0]; r1 = p1.add_run(dfn)
        r1.font.size = Pt(9.5); r1.font.name = 'Calibri'

        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)

    # Save document
    filename = "Proyecto_NutriEvolve_G01_G02_Pediatrico.docx"
    doc.save(filename)
    print(f"Document saved successfully as '{filename}'")

if __name__ == "__main__":
    create_document()
